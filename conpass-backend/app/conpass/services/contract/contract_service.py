import base64
import datetime
import unicodedata

import pdfkit
import tempfile
import traceback
import os
import io
import qrcode
import requests
from qrcode.image.svg import SvgPathImage
from logging import getLogger
from meilisearch import Client, errors
from celery import shared_task
import urllib.parse
from django.utils.html import strip_tags
from rest_framework import status

from django.db import DatabaseError, connection
from django.db.models import Prefetch, OuterRef, Exists, Q, Subquery
from django.utils import html
from django.http import QueryDict
from django.utils.timezone import make_aware
from django.conf import settings

from conpass.models import ContractBody, ContractComment, Contract, User, MetaData, MetaKey, Directory, \
    CorrectionRequest, ContractBodySearch, ConversationComment, Permission, PermissionTarget, LeaseKey
from conpass.models.constants import ContractTypeable, Statusable
from conpass.models.constants.contractmetakeyidable import ContractMetaKeyIdable
from conpass.models.constants.contractstatusable import ContractStatusable
from conpass.services.directory.directory_service import DirectoryService

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import RGBColor, Pt, Mm, Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from PIL import Image
import jaconv

logger = getLogger(__name__)


def _tables_to_markdown(html_body: str) -> str:
    """
    HTML 内の <table> を Markdown テーブル形式に変換する。
    他の HTML タグはそのまま残す（その後 strip_tags で除去する）。
    """
    soup = BeautifulSoup(html_body, "html.parser")
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if not rows:
            table.decompose()
            continue
        md_rows = []
        is_first = True
        for row in rows:
            cells = row.find_all(["th", "td"])
            cell_texts = [
                c.get_text(separator=" ", strip=True).replace("|", "｜")
                for c in cells
            ]
            if not cell_texts:
                continue
            md_rows.append("| " + " | ".join(cell_texts) + " |")
            if is_first:
                md_rows.append("| " + " | ".join(["---"] * len(cell_texts)) + " |")
                is_first = False
        table.replace_with(NavigableString("\n" + "\n".join(md_rows) + "\n"))
    return str(soup)


META_DATA_COMPANY_NAME = [2, 3, 4, 5]  # 会社名（甲乙丙丁）
META_INFO_DATE = [6, 7, 8, 10, 13]  # メタ情報日付項目のID(MetaKey.id)
META_DATA_CONPASS_PERSON = 17  # 担当者名
META_INFO_AMOUNT = [19]  # メタ情報金額項目のID


class ContractService:
    def save_body(self, contract_id: int, version: str, body: str, is_provider: int, user_id: int):
        """
        contractに紐づく本文を作成
        """
        now = make_aware(datetime.datetime.now())
        # <table> を Markdown に変換してから、残りのタグを除去する
        body = _tables_to_markdown(body)
        body = html.strip_tags(body)
        is_ver1st_none = False
        try:
            # versionが1.0の時は既存の1.0のbodyが空かどうかを確認する
            if version == '1.0':
                contractBody = ContractBody.objects.filter(contract_id=contract_id, version=version).first()
                # bodyが空の場合はbodyだけ更新する
                if contractBody and not contractBody.body:
                    contractBody.body = body
                    contractBody.updated_at = now
                    contractBody.updated_by_id = user_id
                    is_ver1st_none = True
            if not is_ver1st_none:
                contractBody = ContractBody(contract_id=contract_id, version=version, body=body, status=ContractBody.Status.ENABLE.value,
                                            created_at=now, created_by_id=user_id,
                                            updated_at=now, updated_by_id=user_id,)
            contractBody.save()
            # 契約書も更新する
            contract = Contract.objects.get(pk=contract_id)
            if is_provider is not None:
                contract.is_provider = True if is_provider else False
            contract.updated_at = now
            contract.updated_by_id = user_id
            contract.save()
        except DatabaseError as e:
            logger.error(f"{e}: {traceback.format_exc()}")
            code = e.args[0]
            if code == 1062:
                raise self.VersionDuplicateError("バージョンが重複しています。")
            else:
                raise e

        # 全検索用モデルとMeilisearchに保存
        try:
            self.save_contract_body_search_task(contractBody, now)
        except Exception as e:
            logger.error(f"contract_body_search error:{e}")
        return True

    def save_contract_body_search_task(self, contract_body: ContractBody, now: datetime):
        """
        全検索用モデルとMeilisearchに保存を行う
        エラーになっても処理を終わらせないようにする
        """
        # ローカルの場合は使用しない
        if settings.RUN_ENV == "local":
            logger.info('runenv is local')
            return

        client = Client(settings.MEILISEARCH_HOST, settings.MEILISEARCH_API_KEY)
        index_name = f"contractbodysearch_{settings.RUN_ENV}_{contract_body.contract.account.id}"
        try:
            # インデックスの存在を確認
            client.get_index(index_name)
        except errors.MeilisearchApiError as e:
            # インデックスが存在しない場合は作成し、設定を適用
            if e.code == 'index_not_found':
                client.create_index(index_name, {'primaryKey': 'contract_body_id'})
                # インデックスオブジェクトを取得して設定を更新
                index = client.index(index_name)
                index.update_settings({
                    'typoTolerance': {
                        'enabled': False
                    },
                    'localizedAttributes': [
                        {
                            'locales': ['jpn'],
                            'attributePatterns': ['*']
                        }
                    ]
                })

        # 保存ようにデコードする
        decoded_body = urllib.parse.unquote(contract_body.body)
        decoded_body = strip_tags(decoded_body)
        # インデックスに保存
        index = client.index(f"contractbodysearch_{settings.RUN_ENV}_{contract_body.contract.account.id}")
        document = {
            'contract_body_id': contract_body.id,
            'contract_id': contract_body.contract.id,
            'search_body': decoded_body
        }
        index.add_documents([document], primary_key='contract_body_id')

        # ContractBodySearchをDBに保存
        contract_body_search = ContractBodySearch.objects.create(
            contract_body=contract_body,
            search_body=decoded_body,
            created_at=now,
            created_by=contract_body.created_by,
            updated_at=now,
            updated_by=contract_body.updated_by
        )
        contract_body_search.save()

    def save_comment(self, contract_id: int, version: str, comment: str, user_id: int):
        now = make_aware(datetime.datetime.now())
        # url encode されているかどうかに関わらずタグは除去する
        comment = html.strip_tags(comment)
        try:
            contractComment = ContractComment(contract_id=contract_id, linked_version=version, comment=comment,
                                              status=ContractComment.Status.ENABLE.value, created_at=now,
                                              created_by_id=user_id, updated_at=now, updated_by_id=user_id)

            contractComment.save()
            # 契約書も更新する
            contract = Contract.objects.get(pk=contract_id)
            contract.updated_at = now
            contract.updated_by_id = user_id
            contract.save()
            return contractComment.id
        except DatabaseError as e:
            logger.error(f"{e}: {traceback.format_exc()}")
            raise e

    def delete_comment(self, comment_id: int, user_id: int):
        # コメントのステータスを更新するロジックをここに追加
        now = make_aware(datetime.datetime.now())
        comment = ContractComment.objects.get(id=comment_id)
        comment.status = Statusable.Status.DISABLE.value
        comment.updated_at = now
        comment.updated_by_id = user_id
        comment.save()

    def save_meta(self, contract_id: int, metadata):
        pass

    def adopt_version(self, contract_id: int, version: str, user_id: int):
        now = make_aware(datetime.datetime.now())
        logger.info(f"version{version}")
        try:
            ContractBody.objects.filter(contract_id=contract_id, is_adopted=True).update(is_adopted=False, updated_at=now, updated_by_id=user_id)

            ContractBody.objects.filter(contract_id=contract_id, version=version).update(is_adopted=True, updated_at=now, updated_by_id=user_id)

            # 契約書も更新する
            contract = Contract.objects.get(pk=contract_id)
            contract.updated_at = now
            contract.updated_by_id = user_id
            contract.save()
        except DatabaseError as e:
            logger.error(f"{e}: {traceback.format_exc()}")
            raise e
        return True

    def get_contract_with_meta_list(self, tags: [str], user: User):
        """
        （契約書とそれに紐づくメタ情報）のリストを取得する
        """
        # ユーザが保有しているディレクトリの権限のリスト取得
        directory_service = DirectoryService()
        visible_directories = directory_service.get_allowed_directories(user, ContractTypeable.ContractType.CONTRACT.value)
        visible_directory_ids = []
        for dir in visible_directories:
            visible_directory_ids.append(dir.id)

        wheres = {
            'account_id': user.account_id,
            'type': ContractTypeable.ContractType.CONTRACT.value,
            'directory_id__in': visible_directory_ids,
        }
        excludes = {
            'status': Contract.Status.DISABLE.value,

        }
        metadata_wheres = {
            'status': MetaData.Status.ENABLE.value,
            'key__type': MetaKey.Type.DEFAULT.value,
            'key__label__in': tags,
            'key__status': MetaKey.Status.ENABLE.value,
        }
        results = []
        try:
            contracts = list(Contract.objects.filter(**wheres).exclude(**excludes).select_related('client').all()
                             .prefetch_related(Prefetch('meta_data_contract',
                                                        queryset=MetaData.objects.filter(**metadata_wheres)
                                                        .select_related('key')
                                                        )))
            for contract in contracts:
                meta = contract.meta_data_contract
                meta_list = []
                if meta:
                    meta_list = list(meta.all())
                if len(meta_list):  # なかったものは除外
                    results.append({
                        'contract': contract,
                        'meta_list': meta_list
                    })
        except Exception as e:
            raise e
        return results

    def create_pdf(self, title: str, body: str, seq=None, dir_name=None):
        """
        契約書PDF作成
        """
        tf = tempfile.NamedTemporaryFile(dir=dir_name)
        tmpname = tf.name + ".pdf"
        absolute_path = ''
        try:
            options = {'title': title}
            if seq is not None:
                # QRコード画像ファイルを生成
                qr = qrcode.QRCode(version=1, box_size=4, border=5)
                qr.add_data('{"id":' + seq + '}')
                qr.make(fit=True)
                img = qr.make_image(fill_color='black', back_color='white', image_factory=SvgPathImage)
                # SVGコードを取得
                buffer = io.BytesIO()
                img.save(buffer)
                svg_code = buffer.getvalue().decode("utf-8")
                # Footer にSVGを埋め込む
                footer_temp = """<!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <title>Footer</title>
                </head>
                <body>
                    <p style="text-align:center;">{svg_code}</p>
                </body>
                </html>
                """
                footer = footer_temp.format(svg_code=svg_code)
                # HTMLファイルを保存
                qr_footer_name = "qr_" + seq + ".html"
                with open(qr_footer_name, "w", encoding="utf-8") as f:
                    f.write(footer)
                absolute_path = os.path.abspath(qr_footer_name)  # Footer ファイルの絶対パスを取得
                # Options を設定
                options['header-right'] = "ID: " + seq
                options['header-font-size'] = 8
                options['header-spacing'] = 5
                options['margin-top'] = "20mm"
                options['margin-bottom'] = "20mm"
                options['footer-html'] = absolute_path
                options['footer-spacing'] = 5
                buffer.close()
            pdfkit.from_string('<meta charset="UTF-8">' + body, tmpname, options=options, css='static/conpass/pdf.css')
            if absolute_path:
                os.remove(absolute_path)  # Header ファイルを削除
            return tmpname
        except Exception as e:
            raise e

    def create_word(self, title: str, body: str, seq=None, dir_name=None):
        """
        契約書Word作成
        """
        tf = tempfile.NamedTemporaryFile(dir=dir_name)
        tmpname = tf.name + ".docx"
        try:
            # bodyからWord文章を生成
            doc = process_html(body)
            if seq is not None:
                # headerを追加
                header = doc.sections[0].header
                header.paragraphs[0].text = "ID: " + seq
                header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # footerを追加
                # QRコード画像ファイルを生成
                qr = qrcode.QRCode(version=1, box_size=4, border=5)
                qr.add_data('{"id":' + seq + '}')
                qr.make(fit=True)
                img = qr.make_image(fill_color='black', back_color='white')
                buffer = io.BytesIO()
                img.save(buffer, format="PNG")
                buffer.seek(0)
                footer = doc.sections[0].footer
                p = footer.paragraphs[0]
                run = p.add_run()
                run.add_picture(buffer, width=Inches(0.5))
                doc.sections[0].footer.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                buffer.close()
            doc.save(tmpname)
            return tmpname
        except Exception as e:
            raise e

    def get_contract_all_list(self, params: dict, user: User):
        """
        主にシステム管理用。ページングも何もないデータ取得
        """
        wheres = {}
        exclude = {
            'status': Contract.Status.DISABLE.value  # いずれにしても無効なデータは出ない
        }
        if params.get('status'):
            wheres['status'] = params.get('status')
        if params.get('nane'):
            wheres['name__contains'] = params.get('name')
        if params.get('account'):
            wheres['account_id'] = params.get('account')
        elif user.type != User.Type.ADMIN.value:
            wheres['account_id'] = user.account_id
        if params.get('client'):
            wheres['client_id'] = params.get('client')
        if params.get('contractType'):
            wheres['type'] = params.get('contractType')
        if params.get('directory'):
            wheres['directory_id'] = params.get('directory')
        if params.get('template'):
            wheres['template_id'] = params.get('template')
        if params.get('origin'):
            wheres['origin_id'] = params.get('origin')

        contract_list = list(Contract.objects.select_related('account', 'client', 'directory', 'template', 'origin')
                             .exclude(**exclude).filter(**wheres).all())
        return contract_list

    def query_search(self, user: User, params: QueryDict):
        account_id = user.account_id

        # query
        type_param = params.get('type')
        contract_type = None
        if type_param == ContractTypeable.ContractType.CONTRACT.value:
            contract_type = ContractTypeable.ContractType.CONTRACT.value
        elif type_param == ContractTypeable.ContractType.TEMPLATE.value:
            contract_type = ContractTypeable.ContractType.TEMPLATE.value

        # 固定の条件(アカウントIDと契約書orテンプレート)
        contract_wheres = {
            'account_id': account_id,
            'type': contract_type,
            'is_garbage': False,
        }
        # ディレクトリ選択時
        directory_service = DirectoryService()
        visible_directories = directory_service.get_allowed_directories(user, contract_type)
        visible_directory_ids = []
        for dir in visible_directories:
            visible_directory_ids.append(dir.id)

        # ゴミ箱選択時はis_garbage == Trueのレコード（階層無視）、それ以外は指定階層のレコードを表示
        if params.get('isGarbage'):
            contract_wheres['is_garbage'] = True
            contract_wheres['directory_id__in'] = visible_directory_ids
        else:
            contract_wheres['is_garbage'] = False
            if params.get('directoryId'):
                target_directories = [params.get('directoryId')]
                # 「全件」の場合、親に紐づく子ディレクトリも検索
                if params.get('showAll'):
                    target_directories.extend(Directory.objects.filter(
                        status=Directory.Status.ENABLE.value,
                        account_id=account_id,
                        parent_id=params.get('directoryId'),
                    ).values_list('id', flat=True).all())
                    for target_dir in target_directories:
                        if target_dir not in visible_directory_ids:
                            target_directories.remove(target_dir)
            else:
                target_directories = visible_directory_ids
            contract_wheres['directory_id__in'] = target_directories

        if params.get('isOpen') is not None:
            contract_wheres['is_open'] = True if params.get('isOpen') else False

        if params.get('isBulkUpload') is not None:
            contract_wheres['bulk_zip_path__isnull'] = False if params.get('isBulkUpload') else True

        # 画面表示が可変の条件
        # 契約書ステータス
        if params.get('status'):
            contract_wheres['status__in'] = [int(x.strip()) for x in params.get('status').split(',')]
        # 契約書作成日
        if params.get('createdAtFrom'):
            created_at_from = datetime.datetime.strptime(params.get('createdAtFrom'), '%Y-%m-%dT%H:%M:%S%z')
            contract_wheres['created_at__gte'] = created_at_from.replace(second=0, microsecond=0)
        if params.get('createdAtTo'):
            created_at_to = datetime.datetime.strptime(params.get('createdAtTo'), '%Y-%m-%dT%H:%M:%S%z')
            # 同じ日時の場合、その分の範囲全てを対象にする（秒とミリ秒を最大値にする）
            contract_wheres['created_at__lte'] = created_at_to.replace(second=59, microsecond=999999)

        # ファイル名
        if params.get('fileName'):
            contract_wheres['file__name__icontains'] = params.get('fileName')

        contracts = Contract.objects.exclude(status=Contract.Status.DISABLE.value).filter(**contract_wheres)
        # デフォルト項目セット
        meta_key_default = MetaKey.objects.filter(type=MetaKey.Type.DEFAULT.value).all()
        for meta in meta_key_default:
            if params.get('default' + str(meta.id)) \
                or params.get('defaultDateFrom' + str(meta.id)) \
                    or params.get('defaultDateTo' + str(meta.id)):
                eq = self._create_exists_query(meta.id,
                                               params.get('default' + str(meta.id)),
                                               params.get('defaultDateFrom' + str(meta.id)),
                                               params.get('defaultDateTo' + str(meta.id))
                                               )
                if eq:
                    contracts = contracts.filter(eq)

            # 金額の場合
            if meta.id in META_INFO_AMOUNT and (params.get('defaultFrom' + str(meta.id)) or params.get('defaultTo' + str(meta.id))):
                from_val = params.get('defaultFrom' + str(meta.id))
                to_val = params.get('defaultTo' + str(meta.id))
                """
                MetaDataのvalueフィールドがCharFieldになっているため、Modelを使用した検索をすると文字列で比較されてしまうため正しく検索ができない。
                valueを数値として比較するために、Modelを使用せずにSQLを直接生成している。
                """
                with connection.cursor() as cursor:
                    try:
                        sql_params = [meta.id, account_id]
                        sql = """
                            SELECT DISTINCT (m.contract_id)
                            FROM conpass_metadata as m
                            INNER JOIN conpass_contract as c
                            ON m.contract_id = c.id
                            WHERE m.key_id = %s
                            AND c.account_id = %s
                        """
                        if from_val:
                            sql += " AND m.value >= %s"
                            sql_params.append(float(from_val))
                        if to_val:
                            sql += " AND m.value <= %s"
                            sql_params.append(float(to_val))
                        cursor.execute(sql, sql_params)
                        metadata = cursor.fetchall()
                        contract_ids = list(map(lambda x: x[0], list(metadata)))
                        contracts = contracts.filter(id__in=contract_ids)
                    except Exception as e:
                        print(e)

        # 自由項目セット
        meta_key_free = MetaKey.objects.filter(type=MetaKey.Type.FREE.value, account_id=account_id).all()
        for meta in meta_key_free:
            if params.get('free' + str(meta.id)):
                contracts = contracts.filter(self._create_exists_query(meta.id,
                                                                       params.get('free' + str(meta.id)),
                                                                       None,
                                                                       None
                                                                       )
                                             )

        # 会社名の横断検索
        if params.get('company'):
            eq = Exists(MetaData.objects.filter(contract_id=OuterRef('id'),
                                                key_id__in=META_DATA_COMPANY_NAME,
                                                value__icontains=params.get('company'),
                                                status=Statusable.Status.ENABLE.value
                                                )
                        )
            contracts = contracts.filter(eq)

        # フルテキスト検索
        if params.get('body') and settings.RUN_ENV != "local":
            # Meilisearchよりcontract_body_idを取得
            contract_ids = self.search_contract_body(params.get('body'), account_id)
            contracts = contracts.filter(id__in=contract_ids)

        # 最後に他テーブルを結合
        contracts = contracts.select_related(
            'account', 'client', 'client__corporate', 'directory', 'template', 'created_by', 'updated_by'
        )
        contracts = contracts.prefetch_related(
            Prefetch('meta_data_contract', queryset=MetaData.objects.select_related('key')),
            Prefetch('correction_request_contract')
        )

        # ソート
        order_by_param = params.get('orderBy')
        if order_by_param:
            meta_key_id_dict = {value.name.lower(): value.value for value in ContractMetaKeyIdable.MetaKeyId}
            # ソート順の指定がある場合はそれを反映する
            order_direction = '-' if order_by_param.endswith('DESC') else ''
            order_key = order_by_param.rstrip(' DESCASC')

            # CorrectionRequestのもの
            if order_key.lower() == 'bpo_correction_response':
                order_subquery = CorrectionRequest.objects \
                    .filter(contract_id=OuterRef('id'), status=Statusable.Status.ENABLE.value) \
                    .order_by('-updated_at') \
                    .values('response')[:1]
                # 新しいフィールドを持たせる
                contracts = contracts.annotate(order_field=Subquery(order_subquery))
                contracts = contracts.order_by(order_direction + 'order_field', order_direction + 'created_at')
            # MetaDetaのもの
            elif order_key.lower() in meta_key_id_dict:
                key_id = meta_key_id_dict[order_key.lower()]
                order_subquery = MetaData.objects \
                    .filter(contract_id=OuterRef('id'), key_id=key_id, status=Statusable.Status.ENABLE.value)
                # 新しいフィールドを持たせる
                if ('date' in order_key.lower() or order_key.lower() == 'cancelnotice') and not order_key.lower() == 'autoupdate':
                    order_subquery = order_subquery.order_by('-date_value')  # 一番新しい日付を一つだけ返す
                    contracts = contracts.annotate(order_field=Subquery(order_subquery.values('date_value')[:1]))
                else:
                    contracts = contracts.annotate(order_field=Subquery(order_subquery.values('value')[:1]))
                contracts = contracts.order_by(order_direction + 'order_field', order_direction + 'created_at')
            # Contractのもの
            else:
                contracts = contracts.order_by(order_direction + order_key, order_direction + 'created_at')
        else:
            contracts = contracts.order_by('-created_at')
        """
        このようなイメージ
        select conpass_contract.id,
            conpass_contract.name,
            conpass_metakey.label,
            conpass_metadata.value
        from conpass_contract
            left join conpass_metadata on conpass_contract.id = conpass_metadata.contract_id
            left join conpass_metakey on conpass_metadata.key_id = conpass_metakey.id
        """
        return contracts

    def _create_exists_query(self, meta_id, value, from_value, to_value):
        """
        メタ情報の検索条件を作成
        """
        # 日付項目の場合(from条件 <= date_value <= toの条件)
        if meta_id in META_INFO_DATE:
            ex_wheres = {
                'contract_id': OuterRef('id'),
                'key_id': meta_id,
                'status': Statusable.Status.ENABLE.value,
            }
            if from_value:
                ex_wheres['date_value__gte'] = datetime.date(*[int(val) for val in from_value.split('-')])
            if to_value:
                ex_wheres['date_value__lte'] = datetime.date(*[int(val) for val in to_value.split('-')])

            eq = Exists(MetaData.objects.filter(**ex_wheres))

        # 担当者はidが入っているが、検索は名前で来る
        elif meta_id == META_DATA_CONPASS_PERSON:
            user_ids = list(User.objects.filter(username__icontains=value).values_list('id', flat=True))
            eq = Exists(MetaData.objects.filter(contract_id=OuterRef('id'),
                                                key_id=meta_id,
                                                value__in=user_ids,
                                                status=Statusable.Status.ENABLE.value
                                                )
                        )
        # その他は文字列部分一致
        else:
            eq = Exists(MetaData.objects.filter(contract_id=OuterRef('id'),
                                                key_id=meta_id,
                                                value__icontains=value,
                                                status=Statusable.Status.ENABLE.value
                                                )
                        )

        return eq

    def search_contract_body(self, search_text, account_id):
        contract_ids = []
        try:
            client = Client(settings.MEILISEARCH_HOST, settings.MEILISEARCH_API_KEY)
            index = client.index(f"contractbodysearch_{settings.RUN_ENV}_{account_id}")
            page = 0
            hits_per_page = 50

            search_keywords = re.split(r'\s+', search_text)
            search_keywords_str = ' '.join(f'"{word}"' for word in search_keywords)
            while True:
                search_results = index.search(search_keywords_str, {
                    'offset': page * hits_per_page,
                    'limit': hits_per_page,
                    'matchingStrategy': 'all',
                    'attributesToRetrieve': ['contract_id'],
                    'locales': ['jpn']
                })

                if not search_results['hits']:
                    break

                contract_ids.extend(hit['contract_id'] for hit in search_results['hits'])
                page += 1
        except errors.MeilisearchApiError as e:
            # インデックスが存在しない場合は空のリストを返す
            if e.code == 'index_not_found':
                return contract_ids
            else:
                raise
        return list(set(contract_ids))

    def get_root_directory(self, contract: Contract) -> Directory:
        current_directory = contract.directory
        while current_directory and current_directory.parent:
            current_directory = current_directory.parent
        return current_directory

    def generate_contract_directory_path(self, contract: Contract):
        current_directory = contract.directory
        parent = []
        while current_directory:
            parent.append(current_directory.name)
            current_directory = current_directory.parent
        return " > ".join([name for name in reversed(parent)])

    def is_allow_to_user(self, contract_id, user: User):
        # ユーザーと契約書のaccount（顧客IDが一致するかどうか
        wheres = {
            'pk': contract_id,
            'account': user.account
        }
        excludes = {
            'status': Contract.Status.DISABLE.value
        }
        try:
            Contract.objects.filter(**wheres).exclude(**excludes).get()
        except Contract.DoesNotExist as e:
            logger.info(e)
            return False
        except Exception as e:
            logger.error(e)
            raise e

        return True

    def get_meta_objects(self, metakey: MetaKey, contract_id: int):
        return MetaData.objects.filter(contract=contract_id, key=metakey.id, status=Statusable.Status.ENABLE.value).order_by('value').all()

    def get_meta_date_value(self, metakey: str, contract_id: int):
        mobjects = self.get_meta_objects(metakey, contract_id)
        return mobjects[0].date_value if mobjects else None

    def get_meta_value(self, metakey: str, contract_id: int):
        mobjects = self.get_meta_objects(metakey, contract_id)
        return mobjects[0].value if mobjects else None

    def expire_contract(self):

        expired_contracts = []

        try:
            # メタデータに「契約終了日」を持つ契約を抽出
            queryset = Contract.objects.prefetch_related('meta_data_contract') \
                .filter(Q(status=Contract.Status.SIGNED.value) | Q(status=Contract.Status.SIGNED_BY_PAPER.value), meta_data_contract__key_id=8).all()
            contracts = list(queryset)

            # 契約終了日と自動更新の情報のmetakeyを抽出
            end_date_metakey = MetaKey.objects.filter(label='contractenddate', status=Statusable.Status.ENABLE.value).get()
            auto_update_metakey = MetaKey.objects.filter(label='autoupdate', status=Statusable.Status.ENABLE.value).get()

            # 「自動更新」が未設定もしくは「自動更新なし」、かつ「契約終了日」が過ぎている場合に「契約満了」にする
            for contract in contracts:
                end_date = self.get_meta_date_value(end_date_metakey, contract.id)
                auto_update = self.get_meta_value(auto_update_metakey, contract.id)
                if (not auto_update or auto_update == "0") and end_date and end_date < datetime.date.today():
                    contract.status = Contract.Status.EXPIRED.value
                    contract.updated_at = make_aware(datetime.datetime.now())
                    contract.save()
                    expired_contracts.append(contract.id)
        except Exception as e:
            raise e
        logger.info(f'contract expiration finished. contract_id:{expired_contracts}')
        return expired_contracts

    def check_contract_permission(self, user: User, contract: Contract):
        uneditable_statuses = [
            ContractStatusable.Status.SIGNED.value,
            ContractStatusable.Status.SIGNED_BY_PAPER.value,
            ContractStatusable.Status.CANCELED.value,
            ContractStatusable.Status.EXPIRED.value,
        ]

        allowed_directories = DirectoryService() \
            .get_allowed_directories(user, contract.type)
        if contract.directory not in allowed_directories:
            return {"message": "対象契約書に対するフォルダの閲覧権限がありません", "status": status.HTTP_400_BAD_REQUEST}
        permission = Permission.objects.filter(
            user=user,
            target=PermissionTarget.Target.DISP_SIGNED_CONTRACT_DETAIL.value
        ).first()

        if permission is not None and not permission.is_allow:
            if contract.status in uneditable_statuses:
                return {"message": "締結済み契約書詳細を閲覧する権限がありません", "status": status.HTTP_400_BAD_REQUEST}

        return {"message": "閲覧可能です", "status": status.HTTP_200_OK}

    def identify_lease_exact_match(self, contract_id: int, text: str):
        # only convert ascii characters and digits
        converted_text = jaconv.z2h(text, kana=False, digit=True, ascii=True)
        normalized_text = unicodedata.normalize("NFKC", converted_text).lower()
        lease_keys=LeaseKey.objects.only("id", "name")
        matched_ids = [
            key.id
            for key in lease_keys
            if unicodedata.normalize("NFKC", key.name) in normalized_text
        ]
        contract = Contract.objects.get(pk=contract_id)
        contract.lease_key.clear()
        contract.is_new_lease=False
        if matched_ids:
            contract.lease_key.set(matched_ids)
            contract.is_new_lease = True
            logger.info(f"New lease identified successfully(exact match)-{contract_id}")
        contract.save()
        return bool(matched_ids)

    def handle_identify_lease_semantic_match(self, contract_id: int, keywords: list):
        contract = Contract.objects.get(pk=contract_id)
        contract.lease_key.clear()
        matching_lease_keys=LeaseKey.objects.filter(name__in=keywords)
        if matching_lease_keys:
            contract.is_new_lease = True
            contract.lease_key.set(matching_lease_keys)
            logger.info(f"New lease identified successfully(semantic match)-{contract_id}")
        contract.save()











    class VersionDuplicateError(Exception):
        pass


def process_html(html_content):
    """
    HTMLコンテンツを処理してWord文書を生成
    """
    html_content = remove_newlines(html_content)  # 改行コードを削除
    soup = BeautifulSoup(html_content, 'html.parser')  # BeautifulSoupで解析
    # brタグを改行コードに置き換える
    for br in soup.find_all("br"):
        br.replace_with("\n")
    doc = Document()
    # 書類のマージンを定める（35mm 30mm 30mm 30mm）
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(35)
        section.bottom_margin = Mm(30)
        section.left_margin = Mm(30)
        section.right_margin = Mm(30)
    for element in soup.find_all():
        process_html_element(element, doc)
    return doc


def remove_newlines(html_content):
    """
    \r\n (Windows), \n (UNIX), \r (Mac) の改行コードを削除
    """
    return re.sub(r'\r\n|\r|\n', '', html_content)


def process_html_element(element, parent=None, inherited_styles={}, inherited_classes=[]):
    """
    HTML要素を解析して各種処理を呼び出す
    """
    try:
        if isinstance(element, NavigableString):
            if parent:
                run = parent.add_run(str(element))
                apply_custom_style(run, inherited_styles, inherited_classes, element)
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            new_styles = {**inherited_styles, **parse_style(element.get('style', ''))}
            p = parent.add_paragraph(style=f'Heading {level}')
            apply_paragraph_styles(p, new_styles)  # 見出しにスタイルを適用
            process_contents(p, element.contents, inherited_styles, inherited_classes)
        elif element.name == 'p':
            new_styles = {**inherited_styles, **parse_style(element.get('style', ''))}
            p = parent.add_paragraph()
            apply_paragraph_styles(p, new_styles)  # Paragraphにスタイルを適用
            process_contents(p, element.contents, inherited_styles, inherited_classes)
        elif element.name == 'span':
            new_classes = inherited_classes + element.get('class', [])
            new_styles = {**inherited_styles, **parse_style(element.get('style', ''))}
            process_contents(parent, element.contents, new_styles, new_classes)
        elif element.name in ['ul', 'ol']:
            style = 'List Bullet' if element.name == 'ul' else 'List Number'
            for li in element.find_all('li', recursive=False):
                p = parent.add_paragraph(style=style)
                for child in li.contents:
                    if child.name == 'p':
                        process_contents(p, child.contents, inherited_styles, inherited_classes)
                    else:
                        process_html_element(child, p, inherited_styles, inherited_classes)
        elif element.name == 'table':
            process_tables(element, parent)
        elif element.name == 'blockquote':
            # blockquoteのスタイルを設定
            blockquote_styles = {**inherited_styles, 'left_indent': 20, 'right_indent': 20}
            for child in element:
                if child.name == 'p':
                    p = parent.add_paragraph()
                    apply_paragraph_styles(p, blockquote_styles)
                    process_contents(p, child.contents, blockquote_styles, inherited_classes)
                else:
                    process_html_element(child, parent, blockquote_styles, inherited_classes)
    except Exception as e:
        logger.error(f"{e}: {traceback.format_exc()}")


def process_contents(p, contents, inherited_styles={}, inherited_classes=[]):
    """
    HTMLコンテンツを処理してWord文書に追加
    """
    for content in contents:
        if isinstance(content, NavigableString):
            # テキストノードの場合、直接テキストを追加
            p.add_run(str(content))
        elif content.name == 'span' and 'strike' in content.get('class', []):
            new_styles = {**inherited_styles, **parse_style(content.get('style', ''))}
            new_classes = inherited_classes + content.get('class', [])
            run = p.add_run(content.text)
            apply_strikethrough(run)
            apply_custom_style(run, new_styles, new_classes, content)
        elif content.name in ['span', 'strong', 'b', 'em', 'u', 'del']:
            new_styles = {**inherited_styles, **parse_style(content.get('style', ''))}
            new_classes = inherited_classes + content.get('class', [])
            # strong, bの場合はboldを適用
            if content.name in ['strong', 'b']:
                new_styles['font-weight'] = 'bold'
            elif content.name == 'em':
                new_styles['font-style'] = 'italic'
            elif content.name == 'u':
                new_styles['text-decoration'] = 'underline'
            elif content.name == 'del':
                new_styles['text-decoration'] = 'line-through'
            for item in content.contents:
                if isinstance(item, NavigableString):
                    run = p.add_run(str(item))
                    apply_custom_style(run, new_styles, new_classes, content)
                else:
                    # 再帰的にprocess_contentsを呼び出し、新しいスタイルとクラスを伝播
                    process_contents(p, [item], new_styles, new_classes)
        elif content.name == 'img':
            img_src = content.get('src')
            img_style = content.get('style', '')
            styles = parse_style(img_style)
            width_style = styles.get("width")
            height_style = styles.get("height")
            width = convert_to_inches(width_style) if width_style else Inches(2)  # 小さめのデフォルト値
            height = convert_to_inches(height_style) if height_style else Inches(2)  # 小さめのデフォルト値
            # Base64エンコードされた画像データを処理
            if is_base64_encoded(img_src):
                decoded_img = get_image_from_base64(img_src)
                image_stream = io.BytesIO(decoded_img)
            else:
                # 外部URLから画像をダウンロード
                response = requests.get(img_src)
                image_stream = io.BytesIO(response.content)
            p.add_run().add_picture(image_stream, width=width, height=height)
        elif content.name == 'br':
            p.add_run('\n')
        elif content.name == 'a':
            url = content.get('href', '')
            text = content.get_text()
            tooltip = content.get('title')
            add_hyperlink(p, url, text, tooltip)
        else:
            p.add_run(content.get_text())


def process_tables(table_tag, parent):
    """
    tableタグを処理してWordのテーブルを追加
    """
    col_count = len(table_tag.find('tr').find_all(['td', 'th']))
    table = parent.add_table(rows=0, cols=col_count)
    for row in table_tag.find_all('tr'):
        row_cells = table.add_row().cells
        for idx, cell in enumerate(row.find_all(['td', 'th'])):
            p = None
            for element in cell.contents:
                if isinstance(element, NavigableString):
                    if not element.strip():
                        continue
                    if p is None:
                        p = row_cells[idx].add_paragraph()
                    p.add_run(str(element))
                else:
                    p = row_cells[idx].add_paragraph()
                    if element.name == 'p':
                        process_contents(p, element.contents)
                    else:
                        process_contents(p, [element])


def parse_style(style_string):
    """
    CSSのstyle属性からスタイルを抽出し、辞書として返す
    """
    styles = {}
    for style in style_string.split(';'):
        if ':' in style:
            key, value = style.split(':', 1)
            styles[key.strip()] = value.strip()
    return styles


def apply_span_styles(run, span, inherited_styles={}):
    """
    spanタグと継承されたスタイルからスタイルを適用
    """
    styles = {**inherited_styles, **parse_style(span.get('style', ''))}
    if 'color' in styles:
        run.font.color.rgb = extract_rgb(styles['color'])


def extract_rgb(color_string):
    """
    CSSの色指定からRGB値を抽出し、RGBColorオブジェクトを返す
    """
    # ff0000のような16進数表記の場合
    if re.match(r'^#[0-9a-fA-F]{6}$', color_string):
        hex_color = color_string.lstrip('#')
        rgb = tuple(int(hex_color[i: i + 2], 16) for i in (0, 2, 4))
        return RGBColor(*rgb)
    # rgb(255, 0, 0)のようなrgb()表記の場合
    elif re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_string):
        color = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)',
                         color_string).groups()
        return RGBColor(int(color[0]), int(color[1]), int(color[2]))
    return None


def apply_paragraph_styles(paragraph, styles):
    """
    paragraphオブジェクトにスタイルを適用
    """
    if 'text-align' in styles:
        if styles['text-align'] == 'center':
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif styles['text-align'] == 'right':
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if 'left_indent' in styles:
        paragraph.paragraph_format.left_indent = Pt(styles['left_indent'])
    if 'right_indent' in styles:
        paragraph.paragraph_format.right_indent = Pt(styles['right_indent'])
    if 'padding-left' in styles:
        padding_left = convert_to_pt(styles['padding-left'])
        paragraph.paragraph_format.left_indent = padding_left


def apply_custom_style(run, styles, custom_classes, content):
    """
    runオブジェクトにスタイルを適用
    """
    if 'font-weight' in styles:
        if styles['font-weight'] == 'bold':
            run.font.bold = True
    if 'font-style' in styles:
        if styles['font-style'] == 'italic':
            run.font.italic = True
    if 'text-decoration' in styles:
        if styles['text-decoration'] == 'underline':
            run.underline = True
        elif styles['text-decoration'] == 'line-through':
            apply_strikethrough(run)
    if 'color' in styles:
        color_rgb = extract_rgb(styles['color'])
        if color_rgb:
            run.font.color.rgb = color_rgb
    if 'font-size' in styles:
        run.font.size = Pt(float(styles['font-size'].rstrip('pt')))
    if 'font-family' in styles:
        run.font.name = styles['font-family']
    if 'text-align' in styles:
        if styles['text-align'] == 'center':
            run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif styles['text-align'] == 'right':
            run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if 'mce-annotation' in custom_classes:
        # contentのdata-mce-annotation-uid属性からuidを取得
        uid = content.get('data-mce-annotation-uid')
        # uidからconversation_commentを取得
        conversation_comment = []
        where = {
            'conversation_id': uid,
            'status': Statusable.Status.ENABLE.value,
        }
        try:
            # 投稿順で全件を取得
            conversation_comment = ConversationComment.objects.filter(**where).order_by('created_at').all()
        except ConversationComment.DoesNotExist:
            logger.error(f"ConversationComment not found. uid: {uid}")
        if conversation_comment:
            # conversation_commentの全件をadd_commentで追加
            for comment in conversation_comment:
                # 表示名（author）があれば使用、なければusernameを使用
                if comment.author is not None:
                    author = comment.author
                elif comment.user.username is not None:
                    author = comment.user.username
                else:
                    author = 'no name'
                run.add_comment(comment.comment, author=author)


def apply_strikethrough(run):
    """
    runオブジェクトに取り消し線を適用
    """
    rPr = run._r.get_or_add_rPr()
    strike = OxmlElement('w:strike')
    strike.set(qn('w:val'), 'true')
    rPr.append(strike)


def convert_to_pt(value):
    """
    CSSのサイズ値（ピクセルなど）をポイントに変換する
    """
    match = re.match(r'(\d+)(\w+)', value)
    if match:
        size, unit = match.groups()
        size = int(size)
        if unit == "px":
            return Pt(size * 0.75)  # 1px ≈ 0.75pt
    return Pt(0)


def convert_to_inches(value):
    """
    CSSのサイズ値をインチに変換する
    """
    if 'in' in value:
        size_in_inches = float(value.replace('in', ''))
        return Inches(size_in_inches)
    elif 'px' in value:
        size_in_px = float(value.replace('px', ''))
        return Inches(size_in_px / 96)  # 1インチ = 96ピクセル
    elif 'cm' in value:
        size_in_cm = float(value.replace('cm', ''))
        return Inches(size_in_cm / 2.54)  # 1インチ = 2.54センチメートル
    return None


def is_base64_encoded(src):
    """
    imgがbase64形式かどうか判定する
    """
    return src.startswith('data:image/')


def get_image_from_base64(src):
    """
    base64形式の画像を取得する
    """
    header, base64_data = src.split(',', 1)
    return base64.b64decode(base64_data)


def add_hyperlink(paragraph, url, text, tooltip=None):
    """
    段落にハイパーリンクを追加するヘルパー関数
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # フォントカラーを青に設定
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # 青色
    rPr.append(color)
    # 下線を設定
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    # ツールチップを設定
    if tooltip:
        rPr_tooltip = OxmlElement('w:tooltip')
        rPr_tooltip.text = tooltip
        rPr.append(rPr_tooltip)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def process_mark_element(run, element):
    """
    mark要素に特定のスタイルを適用する
    """
    # 例：黄色の背景を適用する
    run.font.highlight_color = RGBColor(0xFF, 0xFF, 0x00)
