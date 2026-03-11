#!/usr/bin/env python3
"""
ConPass 5ヵ年事業計画書 最終版 Word文書生成スクリプト
"""

import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("python-docx is not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

import os
from datetime import datetime

# ============================================================
# Helper Functions
# ============================================================

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_key_message_box(doc, text):
    """Add a key message box (highlighted paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add border via XML
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="4" w:color="1F4E79"/>'
        '  <w:left w:val="single" w:sz="4" w:space="4" w:color="1F4E79"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="1F4E79"/>'
        '  <w:right w:val="single" w:sz="4" w:space="4" w:color="1F4E79"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    # Add shading
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="E8F0FE"/>')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def add_formatted_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "1F4E79")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
                run.font.bold = True

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FB")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)

    doc.add_paragraph()  # spacing
    return table

def add_heading_with_number(doc, text, level=1):
    """Add a heading."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_body_text(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_bold_text(doc, bold_part, normal_part=""):
    """Add paragraph with bold prefix."""
    p = doc.add_paragraph()
    run_b = p.add_run(bold_part)
    run_b.bold = True
    run_b.font.size = Pt(10.5)
    if normal_part:
        run_n = p.add_run(normal_part)
        run_n.font.size = Pt(10.5)
    return p

# ============================================================
# Main Document Generation
# ============================================================

def create_document():
    doc = Document()

    # ---- Page Setup ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ---- Header / Footer ----
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "株式会社日本パープル｜ConPass 5ヵ年事業計画書"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = "CONFIDENTIAL"
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ============================================================
    # 表紙
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("ConPass 5ヵ年事業計画書")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("最終版")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()
    doc.add_paragraph()

    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company_para.add_run("株式会社日本パープル")
    run.font.size = Pt(16)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("2026年2月17日")
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf_para.add_run("CONFIDENTIAL")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.bold = True

    conf_note = doc.add_paragraph()
    conf_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf_note.add_run("本書は機密文書です。関係者以外への開示・配布を禁じます。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # ============================================================
    # 目次
    # ============================================================
    add_heading_with_number(doc, "目次", level=1)

    toc_items = [
        ("第1章", "エグゼクティブサマリー", "3"),
        ("第2章", "市場調査", "6"),
        ("第3章", "プロダクト戦略", "12"),
        ("第4章", "財務計画", "20"),
        ("第5章", "GTM・営業戦略", "28"),
        ("第6章", "実行計画とマイルストーン", "36"),
        ("第7章", "リスク管理", "39"),
        ("第8章", "結論と次のステップ", "42"),
        ("付録", "", "44"),
    ]

    for chapter, title, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{chapter}  {title}")
        run.font.size = Pt(11)
        if chapter.startswith("第"):
            run.bold = True
        tab_run = p.add_run(f"\t{page}")
        tab_run.font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # 第1章 エグゼクティブサマリー
    # ============================================================
    add_heading_with_number(doc, "第1章 エグゼクティブサマリー", level=1)

    add_key_message_box(doc,
        "【キーメッセージ】ConPassは「紙も電子も契約書を一元管理できる唯一のクラウドサービス」として、"
        "BPOサービス統合という模倣困難な差別化を武器に、5年間で顧客数318社・ARR 1.59億円の達成を目指す。"
    )

    # 1.1 事業概要
    add_heading_with_number(doc, "1.1 事業概要", level=2)

    add_body_text(doc,
        "ConPassは、株式会社日本パープルが提供する契約管理クラウドサービスである。"
        "紙の契約書も電子契約も一元管理できる唯一のプラットフォームとして、"
        "中小企業（従業員300名以下）の契約管理業務を変革する。"
    )

    add_body_text(doc,
        "日本パープルは年間7,000社・20,000事業所との取引実績を持つ文書保管・廃棄のBPO事業者であり、"
        "このBPO基盤とSaaS技術の融合がConPass最大の競争優位を形成している。"
        "紙契約のスキャン・回収・AI-OCR台帳化（精度98%）・原本保管・機密処理という"
        "End-to-Endフローを提供できるのはConPassだけである。"
    )

    add_body_text(doc,
        "ConPassのビジョンは「商取引インテリジェンス・プラットフォーム」として、"
        "契約管理の効率化を超え、商取引における情報格差を解消し、"
        "企業規模に関係なく健全な取引ができる社会を実現することにある。"
    )

    # 1.2 5ヵ年の主要目標
    add_heading_with_number(doc, "1.2 5ヵ年の主要目標", level=2)

    add_body_text(doc,
        "本計画は55期（2026年）から60期（2031年）の6年間を対象とし、以下の定量目標の達成を目指す。"
    )

    add_formatted_table(doc,
        ["指標", "55期(2026)", "56期(2027)", "57期(2028)", "58期(2029)", "59期(2030)", "60期(2031)"],
        [
            ["顧客数（社）", "79", "118", "159", "203", "254", "318"],
            ["ARR（千円）", "39,368", "58,900", "79,604", "101,609", "127,011", "158,764"],
            ["ARPA（千円/年）", "498", "499", "501", "500", "500", "499"],
            ["前年比成長率", "-", "49.6%", "35.2%", "27.6%", "25.0%", "25.0%"],
        ]
    )

    add_body_text(doc,
        "上記はベースシナリオ（ARPA 50万円維持）の計画値である。"
        "ARPA引き上げ施策（ST比率向上、オプション拡販、ユーザー課金導入）により、"
        "ARPA 100万円を達成した場合、60期ARRは約3.18億円と現計画の2倍に達する可能性がある。"
    )

    # 1.3 投資計画と期待リターン
    add_heading_with_number(doc, "1.3 投資計画と期待リターン", level=2)

    add_body_text(doc,
        "5年間の累計開発投資額は1.25億円（MS1〜MS5）を計画している。"
        "開発投資はプロダクトの成熟に伴い逓減し、粗利による回収は55期から開始される。"
    )

    add_formatted_table(doc,
        ["マイルストーン", "投資額", "期間", "主要内容"],
        [
            ["MS1", "3,000万円", "55期", "AI基盤強化・コア機能開発"],
            ["MS2", "3,000万円", "55-56期", "CLMライフサイクル完成・プッシュ型AI"],
            ["MS3", "3,000万円", "56-57期", "ベンチマークDB公開・外部連携拡張"],
            ["MS4", "2,000万円", "57-58期", "プラットフォーム化・エンタープライズ対応"],
            ["MS5", "1,500万円", "57-58期", "商取引インテリジェンス・プラットフォーム完成"],
        ]
    )

    add_body_text(doc,
        "累計開発投資1.25億円に対し、60期までの累計粗利は約4.5億円に達する見込みであり、"
        "投資回収倍率は3.3倍である。フリーキャッシュフローは59期に黒字転換する。"
    )

    add_bold_text(doc, "資金需要: ", "累計フリーキャッシュフローの最大マイナスは58期末の約1.07億円。"
        "親会社（日本パープル）からの借入1億円と銀行融資の組み合わせで調達する計画である。")

    # 1.4 重要マイルストーン
    add_heading_with_number(doc, "1.4 重要マイルストーン", level=2)

    add_formatted_table(doc,
        ["MS", "完了時期", "テーマ", "ST比率目標", "主要成果"],
        [
            ["MS1", "2026年5月", "AI基盤強化 & ST再生", "20%", "Word取込み、AIリスクチェック、RAG 5ソース、足場業界パック"],
            ["MS2", "2026年9月", "CLMライフサイクル完成", "35%", "プッシュ型AI、電子契約ハブ、3点照合、匿名化パイプライン"],
            ["MS3", "2027年3月", "ベンチマーク公開", "50%", "業界別ベンチマークDB、AIドラフト生成、判例DB統合"],
            ["MS4", "2027年9月", "プラットフォーム化", "60%", "API公開、ISMAP/SOC2取得、エンタープライズプラン"],
            ["MS5", "2028年3月", "商取引インテリジェンス", "70%+", "予測的リスク分析、テンプレートマーケットプレイス、経営ダッシュボード"],
        ]
    )

    add_body_text(doc,
        "各マイルストーンはST（スタンダードプラン）比率の向上を軸に設計されている。"
        "現状10%のST比率を70%以上に引き上げることで、ARPA向上とプロダクト価値の最大化を同時に実現する。"
    )

    # 1.5 市場機会
    add_heading_with_number(doc, "1.5 市場機会", level=2)

    add_body_text(doc,
        "日本国内の契約管理SaaS市場は100〜200億円規模であり、年率20〜25%で成長している。"
        "電子帳簿保存法の完全義務化（2024年1月）、インボイス制度の定着、DX推進施策の加速が追い風となっている。"
    )

    add_body_text(doc,
        "ConPassのターゲット市場（SAM）は中小企業19万社 x ARPA 50万円 = 約950億円/年。"
        "59期時点のSOM（254社、ARR 1.27億円）はSAM浸透率0.13%に過ぎず、成長余地は膨大である。"
    )

    add_body_text(doc,
        "ConPassは「中小企業向け x 低価格 x 導入容易 x BPO統合」というポジションにおいて"
        "唯一の本格プレイヤーであり、競合の空白地帯（ホワイトスペース）を押さえている。"
    )

    doc.add_page_break()

    # ============================================================
    # 第2章 市場調査
    # ============================================================
    add_heading_with_number(doc, "第2章 市場調査", level=1)

    add_key_message_box(doc,
        "【キーメッセージ】日本の契約管理SaaS市場は年率20〜25%で成長中。中小企業のSaaS導入率は10〜15%にとどまり、"
        "膨大な成長余地が存在する。ConPassはBPO統合・価格優位・AI-OCRの3つの競争優位で市場を攻略する。"
    )

    add_heading_with_number(doc, "2.1 市場規模と成長率", level=2)

    add_formatted_table(doc,
        ["セグメント", "2025年推計（億円）", "2030年予測（億円）", "CAGR"],
        [
            ["電子契約サービス全体", "400〜600", "1,000〜1,500", "18〜22%"],
            ["契約管理（CLM）特化", "100〜200", "300〜500", "20〜25%"],
            ["文書管理SaaS全体", "1,500〜2,000", "3,000〜4,000", "15〜18%"],
        ]
    )

    add_body_text(doc,
        "グローバルCLM市場は2025年時点で約25〜30億USD規模、CAGR 15〜17%で2030年には50〜60億USD規模への"
        "拡大が見込まれる。日本市場はアジア太平洋市場の約30〜35%を占め、欧米に比べ3〜5年の遅延があるが、"
        "電帳法義務化を契機に急速にキャッチアップ中である。"
    )

    add_heading_with_number(doc, "2.2 市場の成長ドライバー", level=2)
    add_bullet(doc, "電子帳簿保存法の本格適用（2024年1月完全義務化）")
    add_bullet(doc, "DX推進政策（IT導入補助金2025: SaaS導入に最大450万円）")
    add_bullet(doc, "リモートワーク・ハイブリッドワークの定着")
    add_bullet(doc, "インボイス制度の波及効果（取引書類全般のデジタル管理意識向上）")
    add_bullet(doc, "AI技術の実用化（AI-OCR精度向上、LLMによる契約書レビュー）")

    add_heading_with_number(doc, "2.3 TAM/SAM/SOM", level=2)

    add_formatted_table(doc,
        ["指標", "定義", "金額", "算出根拠"],
        [
            ["TAM", "獲得可能最大市場", "約4,240億円/年", "53万社 x ARPA 80万円"],
            ["SAM", "アプローチ可能市場", "約950億円/年", "19万社 x ARPA 50万円"],
            ["SOM（59期）", "現実的獲得市場", "約1.27億円/年", "254社 x ARPA 50万円"],
        ]
    )

    add_heading_with_number(doc, "2.4 競合分析", level=2)

    add_body_text(doc,
        "ConPassの主要競合はContractsONE（Sansan）、Hubble mini、LegalForce（リーガルオン）である。"
        "いずれも中堅〜大企業向けの高価格帯（月額60,000〜100,000円以上）に位置し、"
        "中小企業向け x 低価格（月額30,000〜50,000円）x BPO統合のポジションはConPassが唯一の本格プレイヤーである。"
    )

    add_formatted_table(doc,
        ["競合", "ARPA（年）", "ターゲット", "ConPassの優位点"],
        [
            ["ContractsONE", "120万円", "中堅〜大企業", "価格30〜50%低、BPO統合"],
            ["Hubble mini", "72万円", "法務部門（中堅以上）", "紙契約対応、中小企業特化"],
            ["LegalForce", "96万円〜", "法務部門（大企業）", "価格優位、導入容易性"],
            ["クラウドサイン", "60〜120万円", "全規模", "契約管理に特化、BPO統合"],
        ]
    )

    add_heading_with_number(doc, "2.5 5層利用者モデル", level=2)

    add_formatted_table(doc,
        ["層", "セグメント", "対象数", "想定ARPA（年）", "市場規模（億円）", "到達時期"],
        [
            ["第1層", "契約の当事者（中小企業）", "約19万社", "50万円", "950", "54期〜"],
            ["第2層", "士業（税理士・弁護士等）", "約10万事業所", "30万円", "300", "55期〜"],
            ["第3層", "業界団体・協会", "約3,000団体", "300万円", "90", "56期〜"],
            ["第4層", "金融機関・保険会社", "約1,500社", "1,000万円", "150", "57期〜"],
            ["第5層", "行政・規制当局", "約1,800自治体", "500万円", "90", "58期〜"],
        ]
    )

    doc.add_page_break()

    # ============================================================
    # 第3章 プロダクト戦略
    # ============================================================
    add_heading_with_number(doc, "第3章 プロダクト戦略", level=1)

    add_key_message_box(doc,
        "【キーメッセージ】ConPassは「商取引インテリジェンス・プラットフォーム」を最終到達点とし、"
        "5つのMoat（ベンチマークDB、ナレッジグラフ、規制準拠認証、BPO統合、業界特化ナレッジ）により"
        "模倣困難な競争優位を段階的に構築する。"
    )

    add_heading_with_number(doc, "3.1 プロダクトビジョン", level=2)

    add_body_text(doc,
        "ConPassが目指すのは単なる契約管理SaaSではない。契約から取引、決済のトレーサビリティを"
        "一気通貫で提供し、日本の商取引における「信用インフラ」となることである。"
    )

    add_body_text(doc,
        "ConPassは3つの核心的価値を社会に届ける。"
    )

    add_bullet(doc, "取引の非対称性解消: ベンチマークDBにより中小企業も業界標準を参照可能にする")
    add_bullet(doc, "商取引の健全性可視化: 契約-証憑-決済の3点照合による自動検証を実現する")
    add_bullet(doc, "経営判断の民主化: 契約データから経営指標を自動生成し、専門家不在でも精度の高い判断を支援する")

    add_heading_with_number(doc, "3.2 5つのMoat戦略", level=2)

    add_formatted_table(doc,
        ["Moat", "構築開始", "競争優位の本質", "模倣困難性"],
        [
            ["1. ベンチマークDB", "MS2", "ネットワーク効果", "データ蓄積に時間を要する"],
            ["2. ナレッジグラフ", "MS1", "知識の構造化（9ソース統合）", "技術的に高難度"],
            ["3. 規制準拠認証", "MS4", "参入障壁（ISMAP/SOC2）", "取得に1年以上＋数千万円"],
            ["4. BPO統合", "MS0", "オペレーション優位", "物理的インフラ構築が困難"],
            ["5. 業界特化ナレッジ", "MS1", "ドメイン知識の蓄積", "専門家協力と時間が必要"],
        ]
    )

    add_heading_with_number(doc, "3.3 プロダクトロードマップ", level=2)

    add_body_text(doc,
        "プロダクト開発は5つのマイルストーン（MS1〜MS5）に沿って段階的に進行する。"
        "各MSは明確なテーマとST比率目標を持ち、プロダクト価値の段階的向上を実現する。"
    )

    add_formatted_table(doc,
        ["MS", "期間", "テーマ", "主要デリバラブル"],
        [
            ["MS1", "〜2026.05", "AI基盤強化 & ST再生",
             "Word取込み+AI差分検出、AIリスクチェック、RAG 5ソース稼働、足場業界パック"],
            ["MS2", "〜2026.09", "CLMライフサイクル完成",
             "AIルーティングWF、電子契約ハブ、プッシュ型監視AI、匿名化パイプライン、3点照合"],
            ["MS3", "〜2027.03", "ベンチマーク公開 & 外部連携",
             "業界別ベンチマーク公開、TDB/TSR連携、判例DB統合、AIドラフト生成"],
            ["MS4", "〜2027.09", "プラットフォーム化",
             "ConPass API公開、ISMAP/SOC2取得、エンタープライズプラン、交渉戦略AI"],
            ["MS5", "〜2028.03", "商取引インテリジェンス",
             "クロス業界ベンチマーク、予測的リスク分析、テンプレートマーケットプレイス"],
        ]
    )

    add_heading_with_number(doc, "3.4 AI・技術戦略", level=2)

    add_body_text(doc,
        "ConPassのAI価値の中核は、9つのナレッジソースを統合したRAGアーキテクチャにある。"
        "現在稼働中の2ソース（社内テンプレート、契約リレーション）を、MS4完了時に9ソース全稼働へ拡充する。"
    )

    add_body_text(doc,
        "AI機能は3つのStageで進化する。Stage 1（MS0-MS1）: 既存文書の理解（AI-OCR・分析）、"
        "Stage 2（MS2-MS3）: 新規文書の生成（AIドラフト・監視）、"
        "Stage 3（MS4-MS5）: 戦略的意思決定支援（交渉AI・予測分析）。"
    )

    add_heading_with_number(doc, "3.5 料金戦略の進化", level=2)

    add_body_text(doc, "ARPA 50万円から100万円への引き上げを3フェーズで実現する。")

    add_formatted_table(doc,
        ["フェーズ", "時期", "主要施策", "想定ARPA"],
        [
            ["Phase 1: ST価値強化", "MS1-MS2", "STにAI基本機能標準搭載、AIプロオプション新設", "55,000円/月"],
            ["Phase 2: ユーザー課金導入", "MS3-MS4", "ユーザー課金（9,000円/人）、エンタープライズプラン", "100,000円/月"],
            ["Phase 3: プラットフォーム課金", "MS5", "ベンチマークDB参照課金、API課金", "150,000円+/月"],
        ]
    )

    doc.add_page_break()

    # ============================================================
    # 第4章 財務計画
    # ============================================================
    add_heading_with_number(doc, "第4章 財務計画", level=1)

    add_key_message_box(doc,
        "【キーメッセージ】55期にブレイクイーブン到達、60期に営業利益黒字化（+16.3%）。"
        "LTV/CAC比率6.7倍、CAC回収期間10.8ヶ月と良好なUnit Economicsを維持する。"
        "累計開発投資1.25億円は60期までの累計粗利4.5億円で3.3倍の回収を見込む。"
    )

    add_heading_with_number(doc, "4.1 売上計画", level=2)

    add_formatted_table(doc,
        ["期", "期首社数", "新規獲得", "解約", "期末社数", "ARR（千円）", "前年比"],
        [
            ["54期(2025)", "21", "27", "3", "45", "14,288", "-"],
            ["55期(2026)", "45", "39", "5", "79", "39,368", "+175.5%"],
            ["56期(2027)", "79", "46", "7", "118", "58,900", "+49.6%"],
            ["57期(2028)", "118", "50", "9", "159", "79,604", "+35.2%"],
            ["58期(2029)", "159", "55", "11", "203", "101,609", "+27.6%"],
            ["59期(2030)", "203", "64", "13", "254", "127,011", "+25.0%"],
            ["60期(2031)", "254", "80", "16", "318", "158,764", "+25.0%"],
        ]
    )

    add_heading_with_number(doc, "4.2 Unit Economics", level=2)

    add_formatted_table(doc,
        ["指標", "値", "業界目安", "判定"],
        [
            ["CAC（顧客獲得コスト）", "45万円", "-", "-"],
            ["LTV（顧客生涯価値）", "301万円", "-", "-"],
            ["LTV/CAC比率", "6.7倍", "3倍以上", "良好"],
            ["CAC回収期間", "10.8ヶ月", "12ヶ月以内", "良好"],
        ]
    )

    add_heading_with_number(doc, "4.3 コスト構造", level=2)

    add_formatted_table(doc,
        ["カテゴリ", "55期", "56期", "57期", "58期", "59期", "60期"],
        [
            ["開発投資", "30,000", "30,000", "30,000", "20,000", "15,000", "10,000"],
            ["人件費", "19,800", "31,500", "44,400", "51,900", "56,400", "57,900"],
            ["インフラ・運用", "3,600", "5,040", "7,140", "9,180", "11,280", "13,320"],
            ["営業・マーケ", "16,158", "19,557", "22,656", "23,976", "26,475", "28,000"],
            ["BPOオペレーション", "4,200", "6,300", "8,400", "10,500", "12,600", "14,700"],
            ["AIアシスタント償却", "2,160", "2,160", "2,160", "2,160", "2,160", "0"],
            ["総コスト", "75,918", "94,557", "114,756", "117,716", "123,915", "123,920"],
        ]
    )

    add_heading_with_number(doc, "4.4 収益性分析", level=2)

    add_formatted_table(doc,
        ["指標", "55期", "56期", "57期", "58期", "59期", "60期"],
        [
            ["ARR（千円）", "39,368", "58,900", "79,604", "101,609", "127,011", "158,764"],
            ["粗利率", "80.2%", "80.7%", "80.5%", "80.6%", "81.2%", "82.3%"],
            ["営業利益率", "-31.9%", "-30.0%", "-29.1%", "-17.8%", "-5.4%", "+16.3%"],
        ]
    )

    add_body_text(doc,
        "SaaS事業として粗利率80%超を安定的に維持し、スケールに伴い逓増傾向にある。"
        "営業利益は60期に黒字化を達成する。55期〜59期は成長投資フェーズとして位置づける。"
    )

    add_heading_with_number(doc, "4.5 キャッシュフロー計画", level=2)

    add_formatted_table(doc,
        ["項目", "55期", "56期", "57期", "58期", "59期", "60期"],
        [
            ["営業CF（千円）", "-490", "1,103", "2,008", "11,553", "26,656", "52,844"],
            ["投資CF（千円）", "-40,800", "-30,000", "-30,000", "-20,000", "-15,000", "-10,000"],
            ["フリーCF（千円）", "-41,290", "-28,897", "-27,992", "-8,447", "+11,656", "+42,844"],
            ["累計フリーCF（千円）", "-41,290", "-70,187", "-98,179", "-106,626", "-94,970", "-52,126"],
        ]
    )

    add_body_text(doc,
        "営業CF黒字化: 56期、フリーCF黒字化: 59期、累計フリーCF回収: 61〜62期見込み。"
        "最大資金需要は58期末の約1.07億円であり、親会社借入と銀行融資で調達する計画。"
    )

    add_heading_with_number(doc, "4.6 3シナリオ分析", level=2)

    add_formatted_table(doc,
        ["シナリオ", "前提", "60期ARR（千円）", "累計損益（千円）"],
        [
            ["楽観", "ARPA75〜100万円、チャーン4%", "270,900", "+86,582"],
            ["ベース", "ARPA50万円、チャーン5.8%", "158,764", "-52,526"],
            ["悲観", "獲得70%、チャーン10%、ARPA45万円", "90,000", "-91,641"],
        ]
    )

    doc.add_page_break()

    # ============================================================
    # 第5章 GTM・営業戦略
    # ============================================================
    add_heading_with_number(doc, "第5章 GTM・営業戦略", level=1)

    add_key_message_box(doc,
        "【キーメッセージ】直販チャネル（展示会・リスティング・ウェビナー）を基盤としつつ、"
        "士業チャネル→業界団体OEM→ITディストリビューターへの段階的なチャネル拡張により、"
        "59期にはパートナー経由MRR比率20%以上を実現する。"
    )

    add_heading_with_number(doc, "5.1 GTM戦略の基本方針", level=2)

    add_bullet(doc, "チャネルレバレッジの最大化: 士業・業界団体・ITディストリビューターを通じた間接チャネルの段階的構築")
    add_bullet(doc, "BPO統合による参入障壁の構築: 純SaaS競合が模倣できない構造的優位性")
    add_bullet(doc, "データネットワーク効果の発動: 契約データの蓄積によるAI精度向上とプラットフォームロックイン")

    add_heading_with_number(doc, "5.2 直販チャネル", level=2)

    add_body_text(doc, "展示会は最大のリード獲得チャネルであり、55期ARR 4,176千円から59期ARR 91,368千円への成長を牽引する。")

    add_formatted_table(doc,
        ["チャネル", "55期ARR（千円）", "59期ARR（千円）", "受注獲得単価", "受注率"],
        [
            ["展示会施策", "4,176", "91,368", "546,296円", "10%→15%"],
            ["ビジネスタンク", "2,808", "28,490", "500,000円", "30%→35%"],
            ["リスティング広告", "1,597", "19,206", "879,121円", "4%→10%"],
            ["既存ウェビナー", "1,022", "27,166", "-", "-"],
            ["SHOKO既存接点", "1,942", "28,836", "-", "-"],
        ]
    )

    add_heading_with_number(doc, "5.3 間接販売チャネル（5層利用者モデル）", level=2)

    add_formatted_table(doc,
        ["フェーズ", "時期", "対象層", "戦略の骨子", "目標"],
        [
            ["Phase 1", "55-56期", "士業（税理士・弁護士）", "士業チャネル確立、顧問先への面展開", "50事務所、100社導入"],
            ["Phase 2", "56-57期", "監査法人", "監査証跡ニーズ、被監査企業への導入", "10社提携、200社導入"],
            ["Phase 3", "57-58期", "業界団体・協会", "OEM提供、不動産→介護→製造", "3団体、500社導入"],
            ["Phase 4", "58-59期", "金融機関・保険会社", "ConPassスコア、融資・保険評価連携", "地銀3行、300社連携"],
        ]
    )

    add_heading_with_number(doc, "5.4 営業組織計画", level=2)

    add_formatted_table(doc,
        ["役割", "55期", "56期", "57期", "58期", "59期"],
        [
            ["営業マネージャー", "1名", "1名", "1名", "2名", "2名"],
            ["フィールドセールス", "2名", "3名", "4名", "5名", "6名"],
            ["インサイドセールス", "1名", "2名", "3名", "3名", "4名"],
            ["パートナーセールス", "0名", "1名", "1名", "2名", "2名"],
            ["CS", "1名", "2名", "3名", "4名", "5名"],
            ["マーケティング", "1名", "1名", "2名", "2名", "2名"],
            ["合計", "6名", "10名", "14名", "18名", "21名"],
        ]
    )

    add_heading_with_number(doc, "5.5 施策別ARR積み上げ", level=2)

    add_formatted_table(doc,
        ["施策", "55期", "56期", "57期", "58期", "59期"],
        [
            ["55期以降ベース", "15,600", "15,600", "15,600", "15,600", "15,600"],
            ["展示会施策", "4,176", "21,384", "44,712", "68,040", "91,368"],
            ["ビジネスタンク", "2,808", "8,078", "14,882", "21,686", "28,490"],
            ["リスティング広告", "1,597", "5,053", "9,771", "14,488", "19,206"],
            ["既存プレーン", "1,022", "5,393", "12,650", "19,908", "27,166"],
            ["SHOKO&ConPass", "156", "444", "732", "1,020", "1,308"],
            ["SHOKO既存接点強化", "1,942", "7,063", "14,321", "21,578", "28,836"],
            ["既存ウェビナー", "1,022", "5,393", "12,650", "19,908", "27,166"],
            ["AIアシスタント", "1,170", "3,315", "5,463", "7,612", "9,761"],
            ["パートナー構築", "840", "4,838", "12,211", "22,349", "35,251"],
            ["Total（千円）", "31,633", "77,861", "144,292", "213,489", "285,450"],
        ]
    )

    add_body_text(doc,
        "注: 施策別積上ARRは各施策の目標値であり、財務計画の基準ARR（顧客数ベース）とは異なる。"
        "財務計画ARRは保守的な前提に基づく計画値であり、施策別積上ARRはその達成に向けた営業活動のKPI目標である。"
    )

    doc.add_page_break()

    # ============================================================
    # 第6章 実行計画とマイルストーン
    # ============================================================
    add_heading_with_number(doc, "第6章 実行計画とマイルストーン", level=1)

    add_key_message_box(doc,
        "【キーメッセージ】5ヵ年の実行計画を4つのフェーズに分け、各フェーズで明確な到達目標を設定する。"
        "組織体制は6名（55期）から21名（59期）へ段階的に拡大し、ガバナンス体制を確立する。"
    )

    add_heading_with_number(doc, "6.1 フェーズ別実行計画", level=2)

    add_heading_with_number(doc, "フェーズ1: 基盤構築期（55期・2026年）", level=3)

    add_body_text(doc, "プロダクト基盤の強化と直販チャネルの確立に集中する。")

    add_bullet(doc, "MS1（〜2026.05）: AI基盤強化、Word取込み、AIリスクチェック、RAG 5ソース稼働")
    add_bullet(doc, "MS2（〜2026.09）: CLMライフサイクル完成、プッシュ型AI転換、電子契約ハブ構築")
    add_bullet(doc, "展示会5回出展、リスティング広告最適化、ウェビナー月2回開催")
    add_bullet(doc, "ダイサン社PoC実施、足場業界パック初版リリース")
    add_bullet(doc, "CS体制1名確立、オンボーディングプログラム体系化")
    add_bullet(doc, "目標: 顧客数79社、ARR 39,368千円")

    add_heading_with_number(doc, "フェーズ2: 拡張期（56-57期・2027-2028年）", level=3)

    add_body_text(doc, "間接チャネルの構築とプラットフォーム化に注力する。")

    add_bullet(doc, "MS3（〜2027.03）: ベンチマークDB公開、AIドラフト生成、判例DB統合")
    add_bullet(doc, "MS4（〜2027.09）: API公開、ISMAP/SOC2取得、エンタープライズプラン")
    add_bullet(doc, "士業チャネル確立: 税理士30事務所、弁護士20事務所のパートナー網構築")
    add_bullet(doc, "バーティカル展開: 足場業界、第2バーティカル（製造業 or IT）")
    add_bullet(doc, "プラン改定によるARPA引き上げ開始（目標ARPA 60〜80万円）")
    add_bullet(doc, "目標: 57期末 顧客数159社、ARR 79,604千円")

    add_heading_with_number(doc, "フェーズ3: スケール期（58-59期・2029-2030年）", level=3)

    add_body_text(doc, "パートナー網の本格活用と大型案件の獲得を推進する。")

    add_bullet(doc, "MS5（〜2028.03）: 商取引インテリジェンス・プラットフォーム完成")
    add_bullet(doc, "監査法人10社提携、業界団体OEM展開（不動産、介護）")
    add_bullet(doc, "ITディストリビューター経由の全国販売開始")
    add_bullet(doc, "ユーザー課金導入、ARPA 100万円達成に向けた施策実行")
    add_bullet(doc, "CS 4〜5名体制、パートナー経由MRR比率20%以上")
    add_bullet(doc, "目標: 59期末 顧客数254社、ARR 127,011千円")

    add_heading_with_number(doc, "フェーズ4: 収穫期（60期・2031年）", level=3)

    add_body_text(doc, "投資回収と次の成長への布石を打つ。")

    add_bullet(doc, "営業利益黒字化（営業利益率+16.3%）の達成")
    add_bullet(doc, "金融機関連携によるConPassスコアの実用化")
    add_bullet(doc, "テンプレートマーケットプレイスの本格稼働")
    add_bullet(doc, "次期中長期計画（61期〜）の策定: ARPA100万円シナリオでのARR 3億円超を視野")
    add_bullet(doc, "目標: 顧客数318社、ARR 158,764千円")

    add_heading_with_number(doc, "6.2 組織体制の進化", level=2)

    add_body_text(doc,
        "組織は事業の成長段階に合わせて段階的に拡大する。55期の6名体制から59期の21名体制へ、"
        "各機能（営業、CS、開発、マーケティング）のバランスを維持しながら増員する。"
    )

    add_formatted_table(doc,
        ["組織機能", "55期（基盤構築）", "57期（拡張期）", "59期（スケール期）"],
        [
            ["営業（FS+IS）", "3名", "7名", "10名"],
            ["パートナーセールス", "0名", "1名", "2名"],
            ["CS", "1名", "3名", "5名"],
            ["マーケティング", "1名", "2名", "2名"],
            ["マネジメント", "1名", "1名", "2名"],
            ["合計", "6名", "14名", "21名"],
        ]
    )

    add_body_text(doc,
        "開発体制は業務委託（BrainStation-23チーム）を中心に運営し、56期から正社員エンジニアの採用を開始する。"
        "59期までに正社員3名＋業務委託の混成チームで開発力を確保する。"
    )

    add_heading_with_number(doc, "6.3 ガバナンス体制", level=2)

    add_bold_text(doc, "意思決定プロセス:")
    add_bullet(doc, "四半期レビュー: 経営層による計画vs実績のレビュー。乖離10%超で計画見直し")
    add_bullet(doc, "月次KPIモニタリング: ARR、顧客数、チャーンレート、ST比率の月次トラッキング")
    add_bullet(doc, "MS完了判定: 各MSのデリバラブル完了を経営判定会議で承認")

    add_bold_text(doc, "報告体制:")
    add_bullet(doc, "週次: 営業パイプラインレビュー、開発進捗報告")
    add_bullet(doc, "月次: KPIダッシュボード報告、CS定例レポート")
    add_bullet(doc, "四半期: 事業計画対比レビュー、投資判断（次MS着手の可否）")
    add_bullet(doc, "半期: 市場環境レビュー、計画のローリング更新")

    add_bold_text(doc, "リスク管理:")
    add_bullet(doc, "リスクレジスターの維持・更新（四半期ごと）")
    add_bullet(doc, "トリガーポイントの設定（顧客獲得達成率80%割れ、チャーン月次1%超など）")
    add_bullet(doc, "エスカレーション基準の明確化と迅速な意思決定体制の確保")

    doc.add_page_break()

    # ============================================================
    # 第7章 リスク管理
    # ============================================================
    add_heading_with_number(doc, "第7章 リスク管理", level=1)

    add_key_message_box(doc,
        "【キーメッセージ】事業計画の実行において、競合リスク、技術リスク、市場リスク、財務リスク、"
        "組織リスクの5カテゴリを管理する。BPO統合という構造的差別化が多くのリスクに対する天然の防御壁となる。"
    )

    add_heading_with_number(doc, "7.1 競合リスク", level=2)

    add_formatted_table(doc,
        ["リスク", "確率", "影響", "対策"],
        [
            ["クラウドサインの契約管理機能強化", "高", "中〜高", "先手としてのパートナーシップ構築、締結→管理の連携確保"],
            ["Sansanの中小企業向け価格改定", "中", "高", "BPO統合という構造的差別化は価格だけでは模倣不可能"],
            ["大手SIerの参入", "中", "中", "中小企業特化のポジションは大手SIerの戦略と競合しにくい"],
            ["海外CLMプレイヤーの日本進出", "低〜中", "中", "大企業向けが中心。中小企業市場への脅威は限定的"],
            ["スタートアップの新規参入", "高", "低〜中", "BPO基盤は新規参入者にとっての参入障壁"],
        ]
    )

    add_heading_with_number(doc, "7.2 技術リスク", level=2)

    add_formatted_table(doc,
        ["リスク", "確率", "影響", "対策"],
        [
            ["AI機能のコモディティ化", "高", "中", "AI機能を差別化の主軸としない。BPO統合が本質的優位"],
            ["LLM APIコスト高騰", "中", "高", "マルチLLM戦略、キャッシュ最適化、OSS LLM検討"],
            ["開発遅延（6ヶ月超）", "低", "中", "スクラム体制、MVP優先、BrainStation-23との継続協業"],
            ["セキュリティインシデント", "低", "極高", "SOC2/ISMAP認証取得（55〜56期で準備開始）"],
        ]
    )

    add_heading_with_number(doc, "7.3 市場リスク", level=2)

    add_formatted_table(doc,
        ["リスク", "確率", "影響", "対策"],
        [
            ["中小企業のIT投資抑制", "中", "高", "ROI訴求強化、IT導入補助金活用、月額3万円は削減対象外"],
            ["DX疲れ・SaaS過多", "中", "中", "「これ1つで完結」というシンプルな価値訴求"],
            ["電帳法等の法改正", "中", "中", "法改正を営業機会に転換する体制整備"],
            ["チャーンレート上昇", "中", "高", "CS先行投資、BPO統合による高いスイッチングコスト"],
        ]
    )

    add_heading_with_number(doc, "7.4 財務リスク", level=2)

    add_formatted_table(doc,
        ["リスク", "確率", "影響", "対策"],
        [
            ["ARPA停滞（50万円横ばい）", "中", "高", "ST比率向上、オプション拡販、ユーザー課金導入"],
            ["顧客獲得ペース鈍化（80%未満）", "低〜中", "高", "チャネル多角化、パートナー施策加速"],
            ["資金ショート", "低", "極高", "親会社借入＋銀行融資で最大1.2億円を確保"],
            ["インフラコスト急騰", "低〜中", "低〜中", "マルチモデル対応、コスト最適化"],
        ]
    )

    add_heading_with_number(doc, "7.5 組織リスク", level=2)

    add_formatted_table(doc,
        ["リスク", "確率", "影響", "対策"],
        [
            ["人材採用の遅延", "中", "高", "採用前倒し、リファラル強化、業務委託活用"],
            ["営業スキルの属人化", "中", "高", "セールスプレイブック整備、ロールプレイング定例化"],
            ["主要人材の離脱", "低", "中", "ドキュメント整備、複数人体制の構築"],
            ["パートナー管理負荷の増大", "中", "中", "パートナーポータル整備、対応の定型化"],
        ]
    )

    add_heading_with_number(doc, "7.6 リスク対応の優先順位", level=2)

    add_body_text(doc, "最も注視すべきリスクは以下の3点である。")

    p = doc.add_paragraph()
    run = p.add_run("1. チャーンレート管理: ")
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run("SaaS事業の成否はNRR（Net Revenue Retention）で決まる。CS先行投資とBPO統合による"
                     "高いスイッチングコストが天然の防御壁だが、能動的なカスタマーサクセス活動への転換が不可欠。")
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run("2. ARPA引き上げ: ")
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run("現行ARPA 50万円のままでは60期ARR 1.59億円にとどまる。ST比率向上とオプション拡販を"
                     "計画通り実行し、ARPA 100万円への道筋を確実にすることが事業成長の鍵。")
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run("3. 競合動向の監視: ")
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run("特にクラウドサインの契約管理機能強化には注意が必要。先手としてのパートナーシップ構築、"
                     "BPO統合の深化により差別化を不可逆的にすることが重要。")
    run.font.size = Pt(10.5)

    doc.add_page_break()

    # ============================================================
    # 第8章 結論と次のステップ
    # ============================================================
    add_heading_with_number(doc, "第8章 結論と次のステップ", level=1)

    add_key_message_box(doc,
        "【キーメッセージ】ConPassは明確な市場機会、独自の競争優位、堅実な財務計画を備えている。"
        "5つのMoat構築とチャネル戦略の段階的実行により、日本の中小企業向け契約管理市場を制覇する。"
    )

    add_heading_with_number(doc, "8.1 計画の要約", level=2)

    add_body_text(doc,
        "ConPass 5ヵ年事業計画書は、以下の4つの柱で構成される。"
    )

    add_bold_text(doc, "1. 明確な市場機会: ",
        "日本国内の契約管理SaaS市場は100〜200億円規模、年率20〜25%成長。"
        "中小企業のSaaS導入率10〜15%という低い浸透率は膨大な成長余地を意味する。"
        "ConPassのSAM 950億円に対し、59期SOM浸透率はわずか0.13%である。")

    add_bold_text(doc, "2. 模倣困難な競争優位: ",
        "BPO統合（日本パープルの7,000社BPO基盤）、AI-OCR台帳化（精度98%）、"
        "中小企業特化の価格設計（月額3〜5万円）の3要素は、純SaaS競合が容易に模倣できない構造的優位性を形成する。"
        "さらに5つのMoat（ベンチマークDB、ナレッジグラフ、規制準拠認証、BPO統合、業界特化ナレッジ）の"
        "段階的構築により、この優位性を不可逆的なものとする。")

    add_bold_text(doc, "3. 堅実な財務計画: ",
        "累計開発投資1.25億円に対し、60期までの累計粗利は4.5億円（回収倍率3.3倍）。"
        "LTV/CAC比率6.7倍、CAC回収期間10.8ヶ月と健全なUnit Economicsを維持する。"
        "60期には営業利益黒字化（+16.3%）を達成する。")

    add_bold_text(doc, "4. 実行可能なGTM戦略: ",
        "展示会を軸とした直販チャネルに加え、士業→業界団体→ITディストリビューターという"
        "段階的なチャネル拡張戦略を策定。各フェーズの成果が次フェーズの推進力となる構造を設計している。")

    add_heading_with_number(doc, "8.2 成功の条件", level=2)

    add_body_text(doc, "本計画の成功には、以下の5つの条件の達成が不可欠である。")

    add_bullet(doc, "ST比率の計画的向上（10%→70%+）: プロダクト価値の継続的強化とアップセル施策の実行")
    add_bullet(doc, "ARPA引き上げ（50万→100万円）: 料金改定、オプション拡販、ユーザー課金の段階的導入")
    add_bullet(doc, "チャーンレートの管理（月次1.5%以下）: CS先行投資、オンボーディング体系化、BPO統合によるスイッチングコスト")
    add_bullet(doc, "パートナーチャネルの構築: 士業50事務所、監査法人10社、ITディストリビューター3社のパートナー網")
    add_bullet(doc, "開発計画の遵守: MS1〜MS5の計画通りの完了、5つのMoatの段階的確立")

    add_heading_with_number(doc, "8.3 次のステップ（直近90日のアクション）", level=2)

    add_body_text(doc, "本計画の承認後、直ちに以下のアクションを実行する。")

    add_formatted_table(doc,
        ["#", "アクション", "担当", "期限", "成果物"],
        [
            ["1", "MS1開発タスクの最終確認と着手", "開発チーム", "2026年2月末", "MS1プロジェクト計画書"],
            ["2", "55期展示会スケジュール確定（5回分）", "マーケティング", "2026年3月中旬", "展示会出展計画"],
            ["3", "ダイサン社PoCキックオフ", "営業チーム", "2026年3月末", "PoC実施計画書"],
            ["4", "CS体制の強化（オンボーディング体系化）", "CS担当", "2026年3月末", "オンボーディングプログラム"],
            ["5", "税理士チャネル開拓の着手", "営業チーム", "2026年4月末", "士業チャネル開拓計画"],
            ["6", "リスティング広告の最適化（LP改善）", "マーケティング", "2026年3月末", "改善後LP・広告設定"],
            ["7", "56期採用計画の策定", "マネジメント", "2026年4月末", "採用計画書"],
            ["8", "四半期レビュー体制の構築", "経営層", "2026年3月末", "KPIダッシュボード"],
        ]
    )

    add_heading_with_number(doc, "8.4 中長期展望（60期以降）", level=2)

    add_body_text(doc,
        "60期以降の中長期展望として、以下の方向性を検討する。"
    )

    add_bullet(doc, "ARR 3億円超の達成: ARPA 100万円シナリオが実現した場合、318社 x 100万円 = 3.18億円のARR")
    add_bullet(doc, "SAM浸透率1%への挑戦: 1,900社の獲得でARR 9.5億円。10年計画での中長期目標")
    add_bullet(doc, "プラットフォームビジネスへの転換: テンプレートマーケットプレイス、API経済圏の確立")
    add_bullet(doc, "ConPassスコアの金融連携: 融資審査・保険評価へのデータ提供によるB2B2B型の収益モデル")
    add_bullet(doc, "VC調達の検討: ARR 2.7億円以上の実績をもってシリーズA調達を検討し、さらなる成長投資を加速")

    add_body_text(doc,
        "ConPassは日本の中小企業に向けた「商取引インテリジェンス・プラットフォーム」として、"
        "契約管理の枠を超え、商取引における情報格差を解消する社会インフラを目指す。"
        "本事業計画書に示した5ヵ年の取り組みは、そのビジョン実現への第一歩である。"
    )

    doc.add_page_break()

    # ============================================================
    # 付録
    # ============================================================
    add_heading_with_number(doc, "付録", level=1)

    add_heading_with_number(doc, "A. 主要前提条件一覧", level=2)

    add_formatted_table(doc,
        ["前提項目", "値", "根拠"],
        [
            ["ARPA（初期）", "500千円/年", "現状実績ベース"],
            ["月次チャーンレート", "0.5%（年間5.8%）", "SaaS業界水準（保守的）"],
            ["粗利率", "80%", "SaaS業界水準"],
            ["ライト/ST比率", "60:40→25:75", "ST移行施策による段階的改善"],
            ["開発投資（MS1-3）", "各30,000千円", "社内見積もりベース"],
            ["業務委託月額", "900千円", "保守100千円+開発800千円"],
            ["展示会出展費", "1,800千円/回", "実績ベース"],
            ["展示会受注率", "10%", "TOKYO開催実績"],
            ["初期費用", "100千円/社", "料金表ベース"],
            ["ソフトウェア償却", "5年定額法", "税法基準"],
        ]
    )

    add_heading_with_number(doc, "B. 用語定義", level=2)

    add_formatted_table(doc,
        ["用語", "定義"],
        [
            ["TAM", "Total Addressable Market。理論上の最大市場規模"],
            ["SAM", "Serviceable Addressable Market。自社がアプローチ可能な市場規模"],
            ["SOM", "Serviceable Obtainable Market。現実的に獲得可能な市場規模"],
            ["ARPA", "Average Revenue Per Account。1顧客あたりの平均年間売上"],
            ["ARR", "Annual Recurring Revenue。年間経常収益"],
            ["MRR", "Monthly Recurring Revenue。月間経常収益"],
            ["NRR", "Net Revenue Retention。既存顧客からの売上維持率"],
            ["CLM", "Contract Lifecycle Management。契約ライフサイクル管理"],
            ["BPO", "Business Process Outsourcing。業務プロセスの外部委託"],
            ["CAC", "Customer Acquisition Cost。顧客獲得コスト"],
            ["LTV", "Lifetime Value。顧客生涯価値"],
            ["ST比率", "スタンダードプランの顧客比率。ARPA向上の重要指標"],
            ["MS", "マイルストーン。プロダクト開発の段階的目標"],
            ["RAG", "Retrieval-Augmented Generation。検索拡張型生成AI"],
        ]
    )

    add_heading_with_number(doc, "C. 市場データソース一覧", level=2)

    add_formatted_table(doc,
        ["データ項目", "参照元", "備考"],
        [
            ["グローバルCLM市場規模", "Grand View Research (2024)", "CAGR 15〜17%"],
            ["グローバルCLM市場予測", "MarketsandMarkets (2024)", "2030年50〜60億USD"],
            ["アジア太平洋CLM市場", "Mordor Intelligence (2024)", "CAGR 20〜25%"],
            ["国内電子契約市場", "矢野経済研究所 (2024)", "2023年度約350億円"],
            ["国内文書管理市場", "ITR (2024)", "SaaS比率増加傾向"],
            ["企業数統計", "総務省「経済センサス」(2021)", "約368万法人"],
            ["中小企業統計", "中小企業庁「中小企業白書」(2024)", "従業員5名以上約53万社"],
            ["クラウド利用率", "総務省「情報通信白書」(2024)", "77.7%"],
        ]
    )

    add_heading_with_number(doc, "D. セクション間整合性に関する注記", level=2)

    add_body_text(doc,
        "本事業計画書の統合にあたり、以下の点について各セクション間の前提条件の違いを明記する。"
    )

    add_bold_text(doc, "ARRの二重構造: ",
        "財務計画（第4章）のARRは「顧客数 x ARPA 50万円」による保守的なベースシナリオであり、"
        "GTM・営業戦略（第5章）の施策別積上ARRは「各施策が計画通り実行された場合の最大ポテンシャル」である。"
        "56期以降、施策別積上ARRが財務計画ARRを上回るのは、施策の目標値が野心的に設定されているためであり、"
        "財務計画はより保守的な前提に基づく。")

    add_bold_text(doc, "ST比率の前提: ",
        "プロダクト戦略（第3章）のST比率はマイルストーン完了時点の瞬間値であり、"
        "財務計画（第4章）のST比率は期間の平均的な構成比として使用している。"
        "両者のアプローチは異なるが、いずれもST比率向上が事業成長の鍵であるという基本認識は共通である。")

    add_bold_text(doc, "開発投資の期別配分: ",
        "プロダクト戦略上のマイルストーン完了時期と、財務計画上の開発投資の期別配分は必ずしも一対一で対応しない。"
        "これは会計処理上の便宜的な期間配分であり、開発投資の総額（MS1-3: 各3,000万円）は整合している。")

    # ---- Final note ----
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- 以上 ---")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本書は2026年2月17日時点の情報に基づく。四半期ごとの実績対比によるローリング更新を推奨する。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    return doc


# ============================================================
# Entry Point
# ============================================================
if __name__ == "__main__":
    output_path = "/Users/hayashi/Desktop/conpass-backend/docs/planning/ConPass_5ヵ年事業計画書_最終版.docx"

    print("ConPass 5ヵ年事業計画書 最終版を生成中...")
    doc = create_document()

    doc.save(output_path)
    print(f"完了: {output_path}")
    print(f"ファイルサイズ: {os.path.getsize(output_path):,} bytes")
